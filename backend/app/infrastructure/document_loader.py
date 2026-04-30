"""Document loading adapter (Unstructured)."""

from __future__ import annotations

import logging
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader
from unstructured.cleaners.core import clean_extra_whitespace
from unstructured.partition.auto import partition

from app.core.config import settings

_MIN_EXTRACTED_CHARS = 200
_HEADING_CATEGORIES = {"Title", "Header", "Heading"}
_TABLE_CATEGORIES = {"Table"}
logger = logging.getLogger(__name__)


def detect_document_type(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix in {".txt", ".md", ".markdown"}:
        return "text"
    if suffix in {".csv", ".tsv", ".xlsx", ".xls"}:
        return "spreadsheet"
    if suffix in {".html", ".htm"}:
        return "html"
    return "unknown"


def _normalize_whitespace(text: str) -> str:
    if not text:
        return ""
    normalized = clean_extra_whitespace(text.replace("\u00a0", " "))
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _is_table_like(text: str) -> bool:
    lines = [line.strip() for line in text.splitlines() if line.strip()][:8]
    return bool(
        lines
        and (
            any("|" in line for line in lines[:6])
            or sum("\t" in line for line in lines) >= 2
            or sum(bool(re.search(r"\S\s{2,}\S", line)) for line in lines) >= 3
        )
    )


def _from_unstructured(file_path: str, *, strategy: str | None = None) -> list[Document]:
    """Load a file via Unstructured and emit one ``Document`` per page.

    Unstructured's ``partition`` returns one element per paragraph / list
    item / heading. If we forwarded each element as its own ``Document``,
    the downstream splitter would never merge them (it operates per-document),
    so a short item like a single resume line would survive as its own tiny
    chunk with almost no embedding signal. Instead, group elements by their
    ``page_number`` (treating ``None`` as a single bucket) and emit one
    cohesive ``Document`` per group. The splitter can then carve real
    ~1800-char parent windows from each page.
    """
    kwargs: dict[str, object] = {"filename": file_path}
    if strategy:
        kwargs["strategy"] = strategy
        kwargs["infer_table_structure"] = True
    logger.info(
        "Unstructured partition start: file=%s strategy=%s",
        file_path,
        strategy or "default",
    )
    try:
        elements = partition(**kwargs)
    except Exception as exc:
        logger.exception(
            "Unstructured partition failed: file=%s strategy=%s error=%s",
            file_path,
            strategy or "default",
            exc,
        )
        raise
    logger.info("Unstructured partition success: element_count=%s", len(elements))
    for idx, element in enumerate(elements, start=1):
        metadata = getattr(element, "metadata", None)
        page_number = getattr(metadata, "page_number", None) if metadata else None
        category = str(getattr(element, "category", "") or "")
        text = _normalize_whitespace(str(element))
        logger.info(
            "Unstructured element[%s]: category=%s page=%s text=%r",
            idx,
            category or "unknown",
            page_number,
            text,
        )

    # Use a list of (key, [texts], [section headings], has_table) to preserve order.
    grouped: list[tuple[int | None, list[str], list[str], bool]] = []
    index_by_key: dict[int | None, int] = {}

    for element in elements:
        text = _normalize_whitespace(str(element))
        if not text:
            continue

        metadata = getattr(element, "metadata", None)
        page_number = getattr(metadata, "page_number", None) if metadata else None
        key = page_number if isinstance(page_number, int) else None
        category = str(getattr(element, "category", "") or "")
        is_heading = category in _HEADING_CATEGORIES
        is_table = category in _TABLE_CATEGORIES or _is_table_like(text)
        element_section = text if is_heading else ""
        rendered_text = text
        if is_heading:
            rendered_text = f"## {text}"
        if is_table:
            rendered_text = f"[TABLE]\n{rendered_text}"

        if key in index_by_key:
            page_no, page_texts, sections, has_table = grouped[index_by_key[key]]
            page_texts.append(rendered_text)
            if element_section and element_section not in sections:
                sections.append(element_section)
            grouped[index_by_key[key]] = (
                page_no,
                page_texts,
                sections,
                has_table or is_table,
            )
        else:
            index_by_key[key] = len(grouped)
            grouped.append(
                (
                    key,
                    [rendered_text],
                    [element_section] if element_section else [],
                    is_table,
                )
            )

    records: list[Document] = []
    for key, texts, sections, has_table in grouped:
        merged = _normalize_whitespace("\n".join(texts))
        if not merged:
            continue
        records.append(
            Document(
                page_content=merged,
                metadata={
                    "page_number": key if isinstance(key, int) else None,
                    "sections": sections,
                    "has_table": has_table,
                },
            )
        )
    return records


def _extract_docx_text(file_path: str) -> list[Document]:
    try:
        doc = DocxDocument(file_path)
    except Exception:
        return []
    texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text and paragraph.text.strip()]
    merged = _normalize_whitespace("\n".join(texts))
    if not merged:
        return []
    return [Document(page_content=merged, metadata={"page_number": None})]


def _extract_pdf_text(file_path: str) -> list[Document]:
    try:
        reader = PdfReader(file_path)
    except Exception:
        return []
    records: list[Document] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = _normalize_whitespace(page.extract_text() or "")
        if not text:
            continue
        records.append(Document(page_content=text, metadata={"page_number": page_index}))
    return records


def load_document(file_path: str) -> list[Document]:
    doc_type = detect_document_type(file_path)
    extraction_tasks: list[tuple[str, callable]] = [("unstructured", lambda: _from_unstructured(file_path))]
    if doc_type == "pdf":
        extraction_tasks.append(("pdf_text", lambda: _extract_pdf_text(file_path)))
        if settings.INGEST_ENABLE_HI_RES_OCR:
            extraction_tasks.append(("ocr", lambda: _from_unstructured(file_path, strategy="hi_res")))
    elif doc_type == "docx":
        extraction_tasks.append(("docx_text", lambda: _extract_docx_text(file_path)))

    candidates: dict[str, list[Document]] = {}
    max_workers = max(1, int(settings.DOCUMENT_EXTRACTION_MAX_WORKERS))
    with ThreadPoolExecutor(max_workers=min(max_workers, len(extraction_tasks))) as pool:
        future_to_name = {pool.submit(task): name for name, task in extraction_tasks}
        for future in as_completed(future_to_name):
            name = future_to_name[future]
            try:
                candidates[name] = future.result()
            except Exception:
                candidates[name] = []

    best = max(candidates.values(), key=lambda docs: sum(len(doc.page_content) for doc in docs), default=[])
    extracted_chars = sum(len(item.page_content) for item in best)
    if extracted_chars >= _MIN_EXTRACTED_CHARS:
        return best

    # If all paths produced limited text, still return the best effort extraction.
    return best
